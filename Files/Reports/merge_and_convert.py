import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def merge_markdown_files(reports_dir):
    """Merges 01_*.md to 12_*.md into merged.md"""
    files = sorted([f for f in os.listdir(reports_dir) if f.endswith('.md') and f.startswith(tuple(f"{i:02d}_" for i in range(1, 13)))])
    
    merged_content = []
    for filename in files:
        filepath = os.path.join(reports_dir, filename)
        with open(filepath, 'r', encoding='utf-8') as infile:
            content = infile.read()
            merged_content.append(content)
            merged_content.append("\n\n---\n\n") # Page separator in Markdown
            
    merged_filepath = os.path.join(reports_dir, "merged.md")
    with open(merged_filepath, 'w', encoding='utf-8') as outfile:
        outfile.write("".join(merged_content[:-1])) # skip last separator
        
    print(f"Successfully merged {len(files)} files into merged.md")
    return merged_filepath

def set_cell_background(cell, fill_hex):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding/margins in dxas (1/20 of a pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def markdown_to_docx(md_path, docx_path):
    """Parses markdown and creates a styled docx file using python-docx."""
    doc = Document()
    
    # Configure page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base Styles Setup
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x2d, 0x37, 0x48) # Slate Gray
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_code_block = False
    code_content = []
    
    in_table = False
    table_rows = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # 1. Handle Code Blocks (Mermaid, ASCII Diagrams, logs)
        if stripped.startswith("```"):
            if in_code_block:
                in_code_block = False
                # Write code block as styled callout box/shaded paragraph
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.right_indent = Inches(0.25)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                
                # Setup borders & shading
                pPr = p._p.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                left_border = OxmlElement('w:left')
                left_border.set(qn('w:val'), 'single')
                left_border.set(qn('w:sz'), '24') # 3pt width
                left_border.set(qn('w:space'), '6')
                left_border.set(qn('w:color'), '3182CE') # Accent Blue
                pBdr.append(left_border)
                pPr.append(pBdr)
                
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F7FAFC"/>') # Light Gray background
                pPr.append(shd)
                
                run = p.add_run("".join(code_content))
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)
                
                code_content = []
            else:
                in_code_block = True
            i += 1
            continue
            
        if in_code_block:
            code_content.append(line)
            i += 1
            continue

        # 2. Handle Tables
        if stripped.startswith("|"):
            in_table = True
            table_rows.append(stripped)
            i += 1
            continue
        elif in_table:
            # Table ended or blank line
            in_table = False
            # Render the table
            render_docx_table(doc, table_rows)
            table_rows = []
            # continue block (do not skip current line if not empty)

        if stripped == "" or stripped == "---":
            # Add page break on markdown separator or handle blank line
            if stripped == "---":
                doc.add_page_break()
            i += 1
            continue
            
        # 3. Headings
        if stripped.startswith("#"):
            h_level = len(stripped) - len(stripped.lstrip("#"))
            title_text = stripped.lstrip("#").strip()
            
            # Clean bold markers inside titles
            title_text = re.sub(r'\*\*(.*?)\*\*', r'\1', title_text)
            
            p = doc.add_heading(level=h_level)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(title_text)
            run.font.name = 'Arial'
            
            if h_level == 1:
                run.font.size = Pt(20)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d) # Deep Navy
            elif h_level == 2:
                run.font.size = Pt(15)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x2b, 0x6c, 0xb0) # Slate Blue
            else:
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x4a, 0x55, 0x68) # Charcoal
            i += 1
            continue
            
        # 4. Handle Lists (Bullet or Numbered)
        if stripped.startswith(("* ", "- ", "• ")):
            list_text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(2.5)
            parse_inline_formatting(p, list_text)
            i += 1
            continue
        elif re.match(r'^\d+\.\s', stripped):
            list_text = re.sub(r'^\d+\.\s', '', stripped).strip()
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(2.5)
            parse_inline_formatting(p, list_text)
            i += 1
            continue
            
        # 5. Standard Paragraphs
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        parse_inline_formatting(p, stripped)
        i += 1

    doc.save(docx_path)
    print(f"Successfully generated styled DOCX: {docx_path}")

def parse_inline_formatting(paragraph, text):
    """Parses markdown bold (**) and italic (*) syntax into docx run styling."""
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for token in tokens:
        if token.startswith('**') and token.endswith('**'):
            run = paragraph.add_run(token[2:-2])
            run.font.bold = True
        elif token.startswith('*') and token.endswith('*'):
            run = paragraph.add_run(token[1:-1])
            run.font.italic = True
        else:
            paragraph.add_run(token)

def render_docx_table(doc, rows):
    """Converts markdown table lines into a styled docx table."""
    # Filter separator lines like |---|---|
    cleaned_rows = []
    for r in rows:
        cells = [c.strip() for c in r.split("|")[1:-1]]
        if all(re.match(r'^:?-+:?$', c) for c in cells if c):
            continue # separator line
        cleaned_rows.append(cells)
        
    if not cleaned_rows:
        return
        
    num_cols = max(len(r) for r in cleaned_rows)
    table = doc.add_table(rows=len(cleaned_rows), cols=num_cols)
    table.autofit = True
    
    # Configure borders and gridlines
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        '<w:left w:val="none"/>'
        '<w:right w:val="none"/>'
        '<w:insideV w:val="none"/>'
        '</w:tblBorders>'
    )
    tblPr.append(tblBorders)

    for row_idx, row_data in enumerate(cleaned_rows):
        row = table.rows[row_idx]
        is_header = (row_idx == 0)
        
        for col_idx, text in enumerate(row_data):
            if col_idx >= len(row.cells):
                break
            cell = row.cells[col_idx]
            
            # Clean bold tags
            clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            
            run = p.add_run(clean_text)
            run.font.size = Pt(9.5)
            
            # Styling cell spacing/paddings
            set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
            
            if is_header:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
                set_cell_background(cell, "2B6CB0") # Accent Blue header
            else:
                if row_idx % 2 == 0:
                    set_cell_background(cell, "F7FAFC") # Zebra striping
                else:
                    set_cell_background(cell, "FFFFFF")

if __name__ == "__main__":
    reports_dir = os.path.dirname(os.path.abspath(__file__))
    merged_md = merge_markdown_files(reports_dir)
    merged_docx = os.path.join(reports_dir, "merged.docx")
    markdown_to_docx(merged_md, merged_docx)
